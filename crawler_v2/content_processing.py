from __future__ import annotations

import hashlib
import os
import pickle
import io
import re
import signal
import subprocess
import sys
import tempfile
import time
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any
from urllib.parse import unquote, urlparse

import pymupdf
from bs4 import BeautifulSoup

from .config import settings
from .legal_metadata import (
    AttachmentLink,
    extract_legal_metadata,
    extract_legal_title_metadata,
    parse_russian_date,
)
from .pdf_ocr import (
    PDF_EXTRACTION_VERSION,
    PDF_FAST_EXTRACTION_VERSION,
    PdfPageText,
    evaluate_text_quality,
    extract_pdf_page_text,
    legal_header_needs_ocr,
)
from .text_repair import repair_ocr_text, technical_filename


WHITESPACE_RE = re.compile(r"\s+")
LINE_SPACES_RE = re.compile(r"[ \t]+")


PDF_OCR_ENGINE_TESSERACT = "tesseract-cli"


class ContentProcessingError(RuntimeError):
    pass


class UnsupportedContentTypeError(
    ContentProcessingError
):
    pass


class EmptyContentError(
    ContentProcessingError
):
    pass


class CorruptDocumentError(
    ContentProcessingError
):
    pass


class CorruptPdfError(
    CorruptDocumentError
):
    pass


class PdfExtractionRuntimeError(
    ContentProcessingError
):
    pass


class PdfExtractionTimeoutError(
    PdfExtractionRuntimeError
):
    pass


@dataclass(frozen=True)
class ExtractedContent:
    title: str
    section: str
    doc_type: str
    content_type: str
    pages: tuple[tuple[int, str], ...]

    content_category: str | None = None
    site_section: str | None = None

    published_at: str | None = None
    document_number: str | None = None
    document_date: str | None = None
    document_kind: str | None = None

    attachments: tuple[
        AttachmentLink,
        ...,
    ] = ()

    extraction_version: str | None = None
    ocr_pages: tuple[int, ...] = ()
    ocr_attempted_pages: tuple[int, ...] = ()
    ocr_rejected_pages: tuple[int, ...] = ()
    ocr_rejection_reasons: tuple[str, ...] = ()
    ocr_timed_out_pages: tuple[int, ...] = ()
    ocr_budget_skipped_pages: tuple[int, ...] = ()
    ocr_deferred_pages: tuple[int, ...] = ()
    unreadable_pages: tuple[int, ...] = ()
    source_page_count: int | None = None

    # Уверенность Tesseract по каждой распознанной странице. Средним по
    # ним заполняется поле `ocr_confidence` в индексе: по нему видно,
    # какие документы стоит перечитать, когда дойдут руки, — эвристическая
    # оценка текста этого не показывает.
    ocr_confidences: tuple[float, ...] = ()


@dataclass(frozen=True)
class TextChunk:
    page_number: int
    chunk_number: int
    text: str


def normalize_text(value: str) -> str:
    lines: list[str] = []

    for raw_line in value.replace(
        "\r\n",
        "\n",
    ).replace(
        "\r",
        "\n",
    ).split("\n"):
        line = LINE_SPACES_RE.sub(
            " ",
            raw_line,
        ).strip()

        if line:
            lines.append(line)

    return "\n".join(lines).strip()


def compact_text(value: str) -> str:
    return WHITESPACE_RE.sub(
        " ",
        value,
    ).strip()


def filename_from_url(url: str) -> str:
    path = unquote(
        urlparse(url).path
    )

    name = PurePosixPath(path).name

    return (
        name
        or "Документ"
    )


SITE_TITLE_PREFIXES = (
    "Администрация города Сочи -",
    "Администрация города Сочи —",
)


COMMON_REMOVE_SELECTORS = (
    "script",
    "style",
    "noscript",
    "template",
    "svg",
    "nav",
    "header",
    "footer",
    "aside",
    "form",

    ".breadcrumb",
    ".breadcrumbs",
    "[class*='breadcrumb']",

    ".main-nav-container",
    ".main-mobile",
    ".footer-menu",

    ".menu",
    ".sidebar",
    ".social",
    ".share",
    ".cookie",
    ".cookie-notification",
    "#hidden-content",

    ".pagination",
    ".print-link",

    ".hidden-lg.hidden-md.hidden-sm",
)


def first_matching_element(
    root: Any,
    selectors: tuple[str, ...],
) -> Any:
    """
    Возвращает первый элемент строго
    по порядку селекторов.

    В отличие от select_one("a, b, c"),
    порядок здесь определяется нами,
    а не положением элемента в HTML.
    """

    for selector in selectors:
        element = root.select_one(
            selector
        )

        if element is not None:
            return element

    return None


def clean_site_title(
    value: str,
) -> str:
    title = compact_text(value)

    for prefix in SITE_TITLE_PREFIXES:
        if title.casefold().startswith(
            prefix.casefold()
        ):
            title = title[
                len(prefix):
            ].strip()

            break

    return title


def extract_page_title(
    soup: BeautifulSoup,
    *,
    final_url: str,
) -> str:
    heading = first_matching_element(
        soup,
        (
            "#ajax-page-content h1",
            "#ajax-page-content h2",
            "#ajax-page-content h3",

            ".page-content h1",
            ".page-content h2",
            ".page-content h3",

            "article h1",
            "article h2",
            "article h3",

            "#maincontent h1",
            "#maincontent h2",
            "#maincontent h3",

            "main h1",
            "main h2",
        ),
    )

    if heading is not None:
        title = clean_site_title(
            heading.get_text(
                " ",
                strip=True,
            )
        )

        if title and len(title) <= 1000:
            return title

    og_title = soup.find(
        "meta",
        attrs={
            "property": "og:title",
        },
    )

    if og_title:
        title = clean_site_title(
            str(
                og_title.get(
                    "content",
                    "",
                )
            )
        )

        if title:
            return title

    if soup.title:
        title = clean_site_title(
            soup.title.get_text(
                " ",
                strip=True,
            )
        )

        if title:
            return title

    return filename_from_url(
        final_url
    )


def fallback_section_from_url(
    final_url: str,
) -> str:
    parsed = urlparse(
        final_url
    )

    path = parsed.path.casefold()
    query = parsed.query.casefold()

    if (
        "normativno-pravovyye-akty"
        in path
        and "element_id=" in query
    ):
        return (
            "Городская власть / "
            "Официальное опубликование "
            "муниципальных правовых актов"
        )

    if "/press-sluzhba/novosti/" in path:
        return "Пресс-служба / Новости"

    if (
        "/press-sluzhba/obyavleniya/"
        in path
    ):
        return "Пресс-служба / Объявления"

    if "/press-sluzhba/" in path:
        return "Пресс-служба"

    return ""


def extract_page_section(
    soup: BeautifulSoup,
    *,
    title: str,
    final_url: str,
) -> str:
    breadcrumb = first_matching_element(
        soup,
        (
            ".breadcrumb",
            ".breadcrumbs",
            ".bx-breadcrumb",
            ".breadcrumb-list",
            ".breadcrumb__list",
            "[itemtype*='BreadcrumbList']",
        ),
    )

    normalized_title = compact_text(
        title
    ).casefold()

    if breadcrumb is not None:
        parts: list[str] = []

        for anchor in breadcrumb.select(
            "a"
        ):
            value = compact_text(
                anchor.get_text(
                    " ",
                    strip=True,
                )
            )

            normalized_value = (
                value.casefold()
            )

            if not value:
                continue

            if normalized_value in {
                "главная",
                "главная страница",
            }:
                continue

            if (
                normalized_title
                and normalized_value
                == normalized_title
            ):
                continue

            if value not in parts:
                parts.append(value)

        if parts:
            section = " / ".join(parts)

            if len(section) <= 300:
                return section

    return fallback_section_from_url(
        final_url
    )


def select_content_container(
    soup: BeautifulSoup,
    *,
    final_url: str,
) -> Any:
    parsed_url = urlparse(
        final_url
    )

    path = parsed_url.path.casefold()
    query = parsed_url.query.casefold()

    is_legal_detail = (
        "normativno-pravovyye-akty"
        in path
        and "element_id=" in query
    )

    if is_legal_detail:
        selectors = (
            "#maincontent",
            ".content#maincontent",
            ".page-content",
            ".content",
            "main",
        )

    elif (
        "/press-sluzhba/novosti/"
        in path
        or "/press-sluzhba/obyavleniya/"
        in path
    ):
        selectors = (
            "#ajax-page-content",
            ".item-detail-content",
            "article.item-detail-text",
            "article",
            ".page-content",
            "#maincontent",
            ".content",
        )

    else:
        selectors = (
            "#ajax-page-content",
            "article.item-detail-text",
            "article",
            "main",
            "[role='main']",
            "#maincontent",
            "#content",
            ".page-content",
            ".content-page",
            ".main-content",
            ".news-detail",
            ".detail",
            ".text",
            ".content",
        )

    container = first_matching_element(
        soup,
        selectors,
    )

    return (
        container
        or soup.body
        or soup
    )


def extract_published_at(
    soup: BeautifulSoup,
) -> str | None:
    """
    Извлекает дату публикации из устойчивого
    элемента сайта:

    <time
        class="detail-content-date"
        datetime="2013-06-21"
    >
    """

    from datetime import date

    element = soup.select_one(
        "time.detail-content-date"
    )

    if element is None:
        # `parse_russian_date` уже импортирован на уровне модуля; локальный
        # импорт только заслонял его тем же именем.
        from crawler_v2.legal_metadata import extract_labeled_value

        return parse_russian_date(
            extract_labeled_value(
                soup.get_text("\n"),
                "Опубликовано на сайте",
            )
        )

    raw_value = element.get(
        "datetime"
    )

    if not raw_value:
        raw_value = element.get_text(
            " ",
            strip=True,
        )

    normalized = str(
        raw_value or ""
    ).strip()

    if not normalized:
        return None

    # datetime может содержать и время,
    # но для сайта достаточно даты.
    candidate = normalized[:10]

    try:
        date.fromisoformat(candidate)
    except ValueError:
        return None

    return candidate


def classify_content_category(
    final_url: str,
) -> tuple[str, str]:
    parsed = urlparse(final_url)

    path = parsed.path.casefold()
    query = parsed.query.casefold()

    if (
        "normativno-pravovyye-akty"
        in path
        and "element_id=" in query
    ):
        return (
            "legal",
            "Нормативные акты",
        )

    if "/press-sluzhba/novosti/" in path:
        return (
            "news",
            "Новости",
        )

    if (
        "/press-sluzhba/obyavleniya/"
        in path
    ):
        return (
            "announcement",
            "Объявления",
        )

    if (
        "/press-sluzhba/mediagalereya"
        in path
        or "/press-sluzhba/mediagalareya"
        in path
    ):
        return (
            "media",
            "Медиагалерея",
        )

    if "/upload/" in path:
        return (
            "file",
            "Файлы",
        )

    if "/press-sluzhba/" in path:
        return (
            "press",
            "Пресс-служба",
        )

    if "/gorodskaya-vlast/" in path:
        return (
            "city_authority",
            "Городская власть",
        )

    if "/zhizn-goroda/" in path:
        return (
            "city_life",
            "Жизнь города",
        )

    if (
        "/afisha-goroda/" in path
        or "/kalendar-meropriyatiy/"
        in path
    ):
        return (
            "events",
            "Афиша города",
        )

    if "/gorod/" in path:
        return (
            "city",
            "О городе",
        )

    return (
        "page",
        "Страница сайта",
    )


def clean_html_container(
    container: Any,
    *,
    final_url: str,
    title: str,
) -> Any:
    """
    Создаёт копию контентного контейнера.

    Исходный BeautifulSoup не изменяется,
    поэтому извлечение даты и других
    метаданных остаётся независимым.
    """

    container_copy = BeautifulSoup(
        str(container),
        "html.parser",
    )

    for selector in COMMON_REMOVE_SELECTORS:
        for element in container_copy.select(
            selector
        ):
            element.decompose()

    path = urlparse(
        final_url
    ).path.casefold()

    is_press_detail = (
        "/press-sluzhba/novosti/" in path
        or (
            "/press-sluzhba/obyavleniya/"
            in path
        )
    )

    if not is_press_detail:
        return container_copy

    press_remove_selectors = (
        ".detail-content-date",
        ".detail-content-small-info",
        ".detail-content-footer",
        ".btn_obs",
        ".modalop",
        ".buttons-rule",
        "a.share",
        ".share",
    )

    for selector in press_remove_selectors:
        for element in container_copy.select(
            selector
        ):
            element.decompose()

    normalized_title = compact_text(
        title
    ).casefold()

    # Заголовок уже хранится в отдельном
    # поле Elasticsearch, поэтому убираем
    # его точный дубль из основного текста.
    if normalized_title:
        for heading in container_copy.select(
            "h1, h2, h3"
        ):
            heading_text = compact_text(
                heading.get_text(
                    " ",
                    strip=True,
                )
            ).casefold()

            if heading_text == normalized_title:
                heading.decompose()
                break

    # Резервная очистка на случай изменения
    # CSS-классов на сайте.
    for unwanted_text in (
        "Перейти к обсуждению",
        "Поделиться",
    ):
        matches = container_copy.find_all(
            string=lambda value: (
                value
                and compact_text(
                    str(value)
                ).casefold()
                == unwanted_text.casefold()
            )
        )

        for match in matches:
            parent = match.parent

            if parent is not None:
                parent.decompose()

    return container_copy


def extract_html(
    body: bytes,
    *,
    content_type: str,
    final_url: str,
) -> ExtractedContent:
    soup = BeautifulSoup(
        body,
        "html.parser",
    )

    title = extract_page_title(
        soup,
        final_url=final_url,
    )

    section = extract_page_section(
        soup,
        title=title,
        final_url=final_url,
    )

    published_at = extract_published_at(
        soup
    )

    (
        content_category,
        site_section,
    ) = classify_content_category(
        final_url
    )

    original_container = (
        select_content_container(
            soup,
            final_url=final_url,
        )
    )

    container = clean_html_container(
        original_container,
        final_url=final_url,
        title=title,
    )

    text = normalize_text(
        container.get_text(
            "\n",
            strip=True,
        )
    )

    if len(text) < settings.min_text_length:
        raise EmptyContentError(
            "В HTML найдено слишком мало текста: "
            f"{len(text)} символов"
        )

    legal_metadata = extract_legal_metadata(
        soup,
        container,
        text=text,
        final_url=final_url,
    )

    return ExtractedContent(
        title=title,
        section=section,
        doc_type="page",
        content_type=(
            content_type
            or "text/html"
        ),
        pages=(
            (
                1,
                text,
            ),
        ),

        content_category=(
            content_category
        ),
        site_section=site_section,

        published_at=published_at,
        document_number=(
            legal_metadata.document_number
        ),
        document_date=(
            legal_metadata.document_date
        ),
        document_kind=(
            legal_metadata.document_kind
        ),
        attachments=(
            legal_metadata.attachments
        ),
    )


PDF_OCR_MODE_FAST = "fast"
PDF_OCR_MODE_BOUNDED = "bounded"
PDF_OCR_MODES = {
    PDF_OCR_MODE_FAST,
    PDF_OCR_MODE_BOUNDED,
}


def _terminate_subprocess(
    process: Any,
    *,
    process_group: bool,
) -> None:
    try:
        running = process.poll() is None
    except Exception:
        running = True

    if not running:
        return

    try:
        if process_group:
            os.killpg(
                process.pid,
                signal.SIGTERM,
            )
        else:
            process.terminate()
    except (ProcessLookupError, OSError):
        pass

    try:
        process.wait(timeout=5)
        return
    except subprocess.TimeoutExpired:
        pass
    except Exception:
        return

    try:
        if process_group:
            os.killpg(
                process.pid,
                signal.SIGKILL,
            )
        else:
            process.kill()
    except (ProcessLookupError, OSError):
        pass

    try:
        process.wait(timeout=5)
    except Exception:
        pass


def _wait_subprocess(
    process: Any,
    *,
    timeout_seconds: int,
    process_group: bool,
) -> int:
    try:
        return int(
            process.wait(
                timeout=timeout_seconds
            )
        )

    except subprocess.TimeoutExpired:
        _terminate_subprocess(
            process,
            process_group=process_group,
        )
        raise

    except BaseException:
        _terminate_subprocess(
            process,
            process_group=process_group,
        )
        raise


def _native_page_fallback(
    native_text: str,
    *,
    minimum_length: int,
    attempted: bool,
    error: str,
) -> PdfPageText:
    normalized = normalize_text(
        native_text
    )
    quality = evaluate_text_quality(
        normalized,
        minimum_length=minimum_length,
    )

    return PdfPageText(
        text=(
            normalized
            if quality.readable
            else ""
        ),
        method=(
            "native"
            if quality.readable
            else "unreadable"
        ),
        quality=quality,
        ocr_attempted=attempted,
        error=error,
    )


def _extract_pdf_page_isolated(
    body: Path,
    *,
    page_index: int,
    native_text: str,
    force_header_ocr: bool,
) -> PdfPageText:
    result_handle = tempfile.NamedTemporaryFile(
        mode="wb",
        prefix="sochi-search-pdf-page-result-",
        suffix=".tmp",
        dir=str(body.parent),
        delete=False,
    )
    result_path = Path(result_handle.name)
    result_handle.close()

    command = [
        sys.executable,
        "-B",
        "-m",
        "crawler_v2.pdf_page_ocr_worker",
        "--input",
        str(body),
        "--output",
        str(result_path),
        "--page-index",
        str(page_index),
        "--minimum-length",
        str(settings.min_text_length),
    ]

    if force_header_ocr:
        command.append(
            "--force-header-ocr"
        )

    process: Any = None

    try:
        process = subprocess.Popen(
            command,
            stdin=subprocess.DEVNULL,
        )

        try:
            return_code = _wait_subprocess(
                process,
                timeout_seconds=(
                    settings.pdf_ocr_page_timeout_seconds
                ),
                process_group=False,
            )
        except subprocess.TimeoutExpired:
            print(
                "PDF_PAGE_TIMEOUT "
                f"page={page_index + 1} "
                "timeout_seconds="
                f"{settings.pdf_ocr_page_timeout_seconds}",
                flush=True,
            )
            return _native_page_fallback(
                native_text,
                minimum_length=(
                    settings.min_text_length
                ),
                attempted=True,
                error=(
                    "timeout: OCR страницы превысил "
                    f"{settings.pdf_ocr_page_timeout_seconds} секунд"
                ),
            )

        if return_code != 0:
            return _native_page_fallback(
                native_text,
                minimum_length=(
                    settings.min_text_length
                ),
                attempted=True,
                error=(
                    "OCR-процесс страницы завершился "
                    f"с кодом {return_code}"
                ),
            )

        try:
            with result_path.open("rb") as stream:
                payload = pickle.load(stream)
        except Exception as exc:
            return _native_page_fallback(
                native_text,
                minimum_length=(
                    settings.min_text_length
                ),
                attempted=True,
                error=(
                    "Не удалось прочитать OCR страницы: "
                    f"{type(exc).__name__}: {exc}"
                ),
            )

        if (
            isinstance(payload, dict)
            and payload.get("status") == "ok"
            and isinstance(
                payload.get("result"),
                PdfPageText,
            )
        ):
            return payload["result"]

        message = (
            str(payload.get("message") or "")
            if isinstance(payload, dict)
            else "неизвестный формат результата"
        )
        return _native_page_fallback(
            native_text,
            minimum_length=(
                settings.min_text_length
            ),
            attempted=True,
            error=(
                "Ошибка OCR страницы: "
                f"{message or 'неизвестная ошибка'}"
            ),
        )

    finally:
        if process is not None:
            _terminate_subprocess(
                process,
                process_group=False,
            )

        result_path.unlink(
            missing_ok=True
        )
        result_path.with_name(
            result_path.name + ".writing"
        ).unlink(
            missing_ok=True
        )



def _ocr_page_with_tesseract(
    page: Any,
    native_text: str,
    *,
    minimum_length: int,
    working_directory: Path,
) -> PdfPageText:
    """
    Распознаёт страницу внешним Tesseract.

    Вызывается уже внутри изолированного процесса документа, поэтому PDF
    открыт один раз, а не заново на каждую страницу. Дочерний Tesseract
    входит в ту же process group и снимается вместе с ней по общему
    таймауту документа.
    """

    from .pdf_ocr_engine import ocr_page, resolve_tesseract

    native = normalize_text(native_text)
    native_quality = evaluate_text_quality(
        native,
        minimum_length=minimum_length,
    )

    try:
        binary = resolve_tesseract(settings.pdf_ocr_binary)
    except Exception as exc:
        return _native_page_fallback(
            native,
            minimum_length=minimum_length,
            attempted=False,
            error=f"Tesseract недоступен: {exc}",
        )

    result = ocr_page(
        page,
        binary=binary,
        tessdata=settings.pdf_ocr_tessdata_path,
        languages=settings.pdf_ocr_languages,
        minimum_dpi=settings.pdf_ocr_minimum_dpi,
        maximum_dpi=settings.pdf_ocr_maximum_dpi,
        pixel_budget=settings.pdf_ocr_pixel_budget,
        psm=settings.pdf_ocr_psm,
        oem=settings.pdf_ocr_oem,
        threads=settings.pdf_ocr_threads,
        timeout_seconds=settings.pdf_ocr_page_timeout_seconds,
        working_directory=working_directory,
        deskew=settings.pdf_ocr_deskew,
        binarize=settings.pdf_ocr_binarize,
        sauvola_k=settings.pdf_ocr_sauvola_k,
        confidence_floor=settings.pdf_ocr_confidence_floor,
        retry_sauvola_k=settings.pdf_ocr_retry_sauvola_k,
        retry_rotations=settings.pdf_ocr_retry_rotations,
        ladder_budget_seconds=(
            settings.pdf_ocr_ladder_budget_seconds
        ),
    )

    if result.blank:
        # Пустой лист: разворот, разделитель, оборот бланка. Тратить на
        # него распознавание незачем.
        return PdfPageText(
            text=native if native_quality.readable else "",
            method="native" if native_quality.readable else "blank",
            quality=native_quality,
            ocr_attempted=False,
            error=None if native_quality.readable else "пустая страница",
        )

    recognized = normalize_text(result.text)
    ocr_quality = evaluate_text_quality(
        recognized,
        minimum_length=minimum_length,
    )

    if recognized and ocr_quality.readable:
        return PdfPageText(
            text=recognized,
            method="ocr",
            quality=ocr_quality,
            ocr_attempted=True,
            rotation=result.rotation,
            error=result.error,
            confidence=result.confidence,
        )

    if native_quality.readable:
        return PdfPageText(
            text=native,
            method="native",
            quality=native_quality,
            ocr_attempted=True,
            error=(
                result.error
                or f"OCR отклонён: {ocr_quality.reason}"
            ),
        )

    details = [
        f"native={native_quality.reason}",
        f"ocr={ocr_quality.reason}:{ocr_quality.score}",
        f"confidence={result.confidence:.0f}",
    ]

    if result.timed_out:
        details.append(
            "timeout: "
            f"{settings.pdf_ocr_page_timeout_seconds} секунд"
        )

    if result.error:
        details.append(str(result.error))

    return PdfPageText(
        text="",
        method="unreadable",
        quality=ocr_quality,
        ocr_attempted=True,
        error=" | ".join(details),
    )


def _extract_pdf_in_process(
    body: bytes | Path,
    *,
    content_type: str,
    final_url: str,
    ocr_mode: str = PDF_OCR_MODE_BOUNDED,
) -> ExtractedContent:
    if ocr_mode not in PDF_OCR_MODES:
        raise ValueError(
            f"Неизвестный режим PDF OCR: {ocr_mode}"
        )

    report_progress = isinstance(
        body,
        Path,
    )
    progress_started = time.monotonic()
    last_progress = progress_started

    try:
        pymupdf.TOOLS.mupdf_display_errors(
            False
        )
        pymupdf.TOOLS.mupdf_display_warnings(
            False
        )
    except AttributeError:
        pass

    document: Any = None

    try:
        if isinstance(body, Path):
            document = pymupdf.open(
                filename=str(body),
                filetype="pdf",
            )
        else:
            document = pymupdf.open(
                stream=body,
                filetype="pdf",
            )

        metadata = document.metadata or {}

        title = compact_text(
            str(
                metadata.get(
                    "title",
                    "",
                )
            )
        )

        if not title:
            title = filename_from_url(
                final_url
            )

        pages: list[tuple[int, str]] = []
        ocr_pages: list[int] = []
        ocr_confidences: list[float] = []
        ocr_attempted_pages: list[int] = []
        ocr_rejected_pages: list[int] = []
        ocr_rejection_reasons: list[str] = []
        ocr_timed_out_pages: list[int] = []
        ocr_budget_skipped_pages: list[int] = []
        ocr_deferred_pages: list[int] = []
        unreadable_pages: list[int] = []
        page_errors = 0
        budget_exhausted_reported = False
        ocr_deadline = (
            progress_started
            + max(
                settings.pdf_ocr_page_timeout_seconds,
                settings.pdf_ocr_document_budget_seconds,
            )
        )

        try:
            page_count = int(document.page_count)
        except Exception as exc:
            raise CorruptPdfError(
                "Повреждена структура страниц PDF: "
                f"{type(exc).__name__}: {exc}"
            ) from exc

        if report_progress:
            print(
                "PDF_PROGRESS "
                f"page=0/{page_count} "
                "stage=opened "
                "elapsed_seconds=0",
                flush=True,
            )

        for page_index in range(
            page_count
        ):
            page_number = page_index + 1

            try:
                page = document.load_page(
                    page_index
                )

            except Exception:
                page_errors += 1
                unreadable_pages.append(
                    page_number
                )
                continue

            try:
                # Склейка переносов нужна и здесь, а не только после OCR:
                # в вёрстке правовых актов строка выравнивается по ширине,
                # и слова рвутся дефисом независимо от того, есть у PDF
                # текстовый слой или нет. «благоустрой-\nство» — два
                # токена в индексе, и ни один не найдётся по запросу
                # «благоустройство».
                #
                # Письменность здесь намеренно не чинится: на этом тексте
                # принимается решение, запускать ли OCR, и подмена букв
                # сдвинула бы оценку качества слоя.
                native_text = normalize_text(
                    repair_ocr_text(
                        page.get_text(
                            "text",
                            sort=True,
                        ),
                        fix_mixed_script=False,
                        fix_document_numbers=False,
                    )
                )

            except Exception:
                native_text = ""

            force_header_ocr = (
                page_index == 0
                and legal_header_needs_ocr(
                    native_text
                )
            )

            page_text = _native_page_fallback(
                native_text,
                minimum_length=(
                    settings.min_text_length
                ),
                attempted=False,
                error="",
            )
            needs_ocr = (
                not page_text.quality.readable
                or force_header_ocr
            )
            within_attempt_limit = (
                len(ocr_attempted_pages)
                < settings.pdf_ocr_max_pages
            )

            # Первый проход по корпусу распознаёт ограниченное число
            # страниц документа: этого хватает на вид акта, номер, дату и
            # начало текста, и весь корпус становится искомым за часы, а не
            # за недели. Остальное дочитывает следующий, глубокий проход.
            #
            # Считаются именно попытки распознавания, а не номера страниц.
            # Ограничение по номеру (`page_index < N`) выглядит тем же
            # самым, но ведёт себя иначе на смешанном документе: у акта с
            # текстовым слоем на первых листах и отсканированным
            # приложением с третьего распознавать в окне просто нечего,
            # а всё, что нуждается в OCR, лежит за его границей и
            # откладывается навсегда. На рабочем сервере это дало
            # «OCR страниц: 0, отложено в фон: 35».
            within_coverage = (
                settings.pdf_ocr_coverage_pages <= 0
                or len(ocr_attempted_pages)
                < settings.pdf_ocr_coverage_pages
            )
            within_time_budget = (
                time.monotonic()
                < ocr_deadline
            )
            can_attempt_ocr = (
                ocr_mode
                == PDF_OCR_MODE_BOUNDED
                and settings.pdf_ocr_enabled
                and within_attempt_limit
                and within_time_budget
                and within_coverage
            )

            if needs_ocr and can_attempt_ocr:
                if (
                    settings.pdf_ocr_engine
                    == PDF_OCR_ENGINE_TESSERACT
                ):
                    # Внешний Tesseract: документ уже открыт, страница
                    # рендерится в собственном разрешении скана,
                    # бинаризуется и выравнивается по перекосу.
                    page_text = _ocr_page_with_tesseract(
                        page,
                        native_text,
                        minimum_length=(
                            settings.min_text_length
                        ),
                        working_directory=(
                            body.parent
                            if isinstance(body, Path)
                            else settings.download_temp_dir
                        ),
                    )

                elif isinstance(body, Path):
                    page_text = (
                        _extract_pdf_page_isolated(
                            body,
                            page_index=page_index,
                            native_text=native_text,
                            force_header_ocr=(
                                force_header_ocr
                            ),
                        )
                    )
                else:
                    page_text = extract_pdf_page_text(
                        page,
                        native_text,
                        tessdata=(
                            settings.pdf_ocr_tessdata_path
                        ),
                        languages=(
                            settings.pdf_ocr_languages
                        ),
                        dpi=settings.pdf_ocr_dpi,
                        rotations=(
                            settings.pdf_ocr_rotations
                        ),
                        minimum_length=(
                            settings.min_text_length
                        ),
                        ocr_enabled=True,
                        force_ocr=force_header_ocr,
                    )

            elif needs_ocr:
                if ocr_mode == PDF_OCR_MODE_FAST:
                    ocr_deferred_pages.append(
                        page_number
                    )
                elif not settings.pdf_ocr_enabled:
                    ocr_deferred_pages.append(
                        page_number
                    )
                elif not within_coverage:
                    ocr_deferred_pages.append(
                        page_number
                    )
                else:
                    ocr_budget_skipped_pages.append(
                        page_number
                    )

                    if not budget_exhausted_reported:
                        print(
                            "PDF_BUDGET_EXHAUSTED "
                            f"page={page_number}/{page_count} "
                            "attempted="
                            f"{len(ocr_attempted_pages)} "
                            "budget_seconds="
                            f"{settings.pdf_ocr_document_budget_seconds}",
                            flush=True,
                        )
                        budget_exhausted_reported = True

            if page_text.ocr_attempted:
                ocr_attempted_pages.append(
                    page_number
                )

                if page_text.method != "ocr":
                    ocr_rejected_pages.append(
                        page_number
                    )
                    ocr_rejection_reasons.append(
                        "стр. "
                        f"{page_number}: "
                        + (
                            page_text.error
                            or "результат OCR отклонён"
                        )
                    )

                if str(
                    page_text.error or ""
                ).startswith("timeout:"):
                    ocr_timed_out_pages.append(
                        page_number
                    )

            if page_text.method == "ocr":
                ocr_pages.append(
                    page_number
                )

                if page_text.confidence > 0:
                    ocr_confidences.append(
                        page_text.confidence
                    )

            if page_text.text:
                pages.append(
                    (
                        page_number,
                        normalize_text(
                            page_text.text
                        ),
                    )
                )

            else:
                unreadable_pages.append(
                    page_number
                )

            if report_progress:
                progress_now = time.monotonic()
                should_report = (
                    page_number == 1
                    or page_number == page_count
                    or page_number % 10 == 0
                    or progress_now - last_progress >= 30
                )

                if should_report:
                    print(
                        "PDF_PROGRESS "
                        f"page={page_number}/{page_count} "
                        f"readable={len(pages)} "
                        "ocr_attempted="
                        f"{len(ocr_attempted_pages)} "
                        f"ocr_used={len(ocr_pages)} "
                        "ocr_timeouts="
                        f"{len(ocr_timed_out_pages)} "
                        "ocr_skipped="
                        f"{len(ocr_budget_skipped_pages)} "
                        "elapsed_seconds="
                        f"{int(progress_now - progress_started)}",
                        flush=True,
                    )
                    last_progress = progress_now

        total_length = sum(
            len(text)
            for _, text in pages
        )

        if not pages and page_errors >= page_count:
            raise CorruptPdfError(
                "Не удалось прочитать страницы PDF: "
                f"ошибок {page_errors} из {page_count}"
            )

        if total_length < settings.min_text_length:
            # Не сохраняем битый глифовый слой. Метаданные оставляют PDF
            # доступным по заголовку, а прежние мусорные chunks удаляются
            # после успешной переиндексации новой версией extraction.
            fallback_text = (
                title
                or filename_from_url(
                    final_url
                )
            )

            pages = [
                (
                    1,
                    fallback_text,
                )
            ]

        first_page_text = next(
            (
                text
                for page_number, text in pages
                if page_number == 1
            ),
            "",
        )
        legal_metadata = (
            extract_legal_title_metadata(
                first_page_text
            )
        )

        return ExtractedContent(
            title=title,
            section="PDF-документ",
            doc_type="pdf",
            content_type=(
                content_type
                or "application/pdf"
            ),
            pages=tuple(pages),
            document_number=(
                legal_metadata.document_number
            ),
            document_date=(
                legal_metadata.document_date
            ),
            document_kind=(
                legal_metadata.document_kind
            ),
            extraction_version=(
                PDF_FAST_EXTRACTION_VERSION
                if ocr_mode == PDF_OCR_MODE_FAST
                else PDF_EXTRACTION_VERSION
            ),
            ocr_pages=tuple(
                ocr_pages
            ),
            ocr_confidences=tuple(
                ocr_confidences
            ),
            ocr_attempted_pages=tuple(
                ocr_attempted_pages
            ),
            ocr_rejected_pages=tuple(
                ocr_rejected_pages
            ),
            ocr_rejection_reasons=tuple(
                ocr_rejection_reasons
            ),
            ocr_timed_out_pages=tuple(
                ocr_timed_out_pages
            ),
            ocr_budget_skipped_pages=tuple(
                ocr_budget_skipped_pages
            ),
            ocr_deferred_pages=tuple(
                ocr_deferred_pages
            ),
            unreadable_pages=tuple(
                sorted(
                    set(unreadable_pages)
                )
            ),
            source_page_count=page_count,
        )

    except EmptyContentError:
        raise

    except CorruptPdfError:
        raise

    except Exception as exc:
        raise CorruptPdfError(
            f"Ошибка обработки PDF: {exc}"
        ) from exc

    finally:
        if document is not None:
            document.close()


def extract_pdf(
    body: bytes | Path,
    *,
    content_type: str,
    final_url: str,
    ocr_mode: str = PDF_OCR_MODE_BOUNDED,
) -> ExtractedContent:
    if not isinstance(body, Path):
        return _extract_pdf_in_process(
            body,
            content_type=content_type,
            final_url=final_url,
            ocr_mode=ocr_mode,
        )

    result_handle = tempfile.NamedTemporaryFile(
        mode="wb",
        prefix="sochi-search-pdf-result-",
        suffix=".tmp",
        dir=str(body.parent),
        delete=False,
    )
    result_path = Path(result_handle.name)
    result_handle.close()

    command = [
        sys.executable,
        "-B",
        "-m",
        "crawler_v2.pdf_extraction_worker",
        "--input",
        str(body),
        "--output",
        str(result_path),
        "--content-type",
        content_type,
        "--final-url",
        final_url,
        "--ocr-mode",
        ocr_mode,
    ]

    process: Any = None
    processing_timeout_seconds = (
        settings.pdf_fast_processing_timeout_seconds
        if ocr_mode == PDF_OCR_MODE_FAST
        else settings.pdf_processing_timeout_seconds
    )

    try:
        print(
            "PDF_PROCESS "
            f"mode={ocr_mode} "
            f"timeout_seconds="
            f"{processing_timeout_seconds} "
            "page_timeout_seconds="
            f"{settings.pdf_ocr_page_timeout_seconds} "
            "document_budget_seconds="
            f"{settings.pdf_ocr_document_budget_seconds}",
            flush=True,
        )

        try:
            process = subprocess.Popen(
                command,
                stdin=subprocess.DEVNULL,
                start_new_session=True,
            )
            return_code = _wait_subprocess(
                process,
                timeout_seconds=(
                    processing_timeout_seconds
                ),
                process_group=True,
            )
        except subprocess.TimeoutExpired as exc:
            if ocr_mode == PDF_OCR_MODE_FAST:
                fallback_title = filename_from_url(
                    final_url
                )
                print(
                    "PDF_FAST_FALLBACK "
                    "reason=timeout "
                    f"timeout_seconds={processing_timeout_seconds}",
                    flush=True,
                )
                return ExtractedContent(
                    title=fallback_title,
                    section="PDF-документ",
                    doc_type="pdf",
                    content_type=(
                        content_type
                        or "application/pdf"
                    ),
                    pages=((1, fallback_title),),
                    extraction_version=(
                        PDF_FAST_EXTRACTION_VERSION
                    ),
                    ocr_deferred_pages=(1,),
                    unreadable_pages=(1,),
                )

            raise PdfExtractionTimeoutError(
                "Обработка PDF превысила безопасный предел "
                f"{processing_timeout_seconds} секунд"
            ) from exc

        if return_code != 0:
            raise PdfExtractionRuntimeError(
                "Изолированный обработчик PDF завершился "
                f"с кодом {return_code}"
            )

        try:
            with result_path.open("rb") as stream:
                payload = pickle.load(stream)
        except Exception as exc:
            raise PdfExtractionRuntimeError(
                "Не удалось прочитать результат "
                f"обработки PDF: {type(exc).__name__}: {exc}"
            ) from exc

        if not isinstance(payload, dict):
            raise PdfExtractionRuntimeError(
                "Обработчик PDF вернул неизвестный формат"
            )

        status = str(payload.get("status") or "")

        if status == "ok":
            content = payload.get("content")

            if not isinstance(content, ExtractedContent):
                raise PdfExtractionRuntimeError(
                    "Обработчик PDF не вернул ExtractedContent"
                )

            return content

        message = str(
            payload.get("message")
            or "Неизвестная ошибка обработки PDF"
        )
        error_kind = str(
            payload.get("error_kind")
            or "runtime"
        )

        if error_kind == "corrupt":
            raise CorruptPdfError(message)

        if error_kind == "empty":
            raise EmptyContentError(message)

        raise PdfExtractionRuntimeError(message)

    finally:
        if process is not None:
            _terminate_subprocess(
                process,
                process_group=True,
            )

        result_path.unlink(
            missing_ok=True
        )
        result_path.with_name(
            result_path.name + ".writing"
        ).unlink(
            missing_ok=True
        )


def document_title_from_url(final_url: str) -> str:
    """Имя файла, если оно что-то значит; иначе пусто."""

    name = filename_from_url(final_url)

    return "" if technical_filename(name) else name


def describe_document(
    text: str,
    final_url: str,
    *,
    fallback_title: str = "",
) -> tuple[str, Any]:
    """
    Заголовок и реквизиты по тексту самого документа.

    Раньше это делалось только для PDF: `extract_legal_title_metadata`
    вызывался ровно в одном месте — внутри разбора PDF. Офисные файлы
    реквизитов из своего текста не получали никогда, только по
    наследству от родительской HTML-карточки. У файла, найденного прямо
    в `/upload/` без родителя, реквизитов не появлялось вовсе: в выдаче
    карточка с хешем вместо названия, без даты и без номера — при том
    что первая строка внутри документа начинается с «ПОСТАНОВЛЕНИЕ
    администрации города Сочи от … № …».

    Разбор один и тот же для всех форматов: текст есть текст.
    """

    metadata = extract_legal_title_metadata(text)

    title = fallback_title or document_title_from_url(final_url)

    if not title:
        title = legal_title_from_text(text)

    return title, metadata


def legal_title_from_text(text: str) -> str:
    """
    Первая содержательная строка документа как название.

    Берётся начало от вида акта, если он найден: «Постановление
    администрации города Сочи от 14 марта 2024 года № 512-р». Если вида
    акта нет — просто первая непустая строка разумной длины.
    """

    normalized = compact_text(text)

    if not normalized:
        return ""

    match = LEGAL_KIND_IN_TEXT_RE.search(normalized[:1200])
    start = match.start() if match else 0
    candidate = normalized[start:start + 300]

    for separator in ("»", ". ", "; "):
        position = candidate.find(separator, 40)

        if position > 0:
            candidate = candidate[:position + len(separator)]
            break

    candidate = compact_text(candidate).strip(" -–—:;,")

    return candidate[:240]


LEGAL_KIND_IN_TEXT_RE = re.compile(
    (
        r"\b(?:постановление|распоряжение|решение|приказ|"
        r"положение|указ|извещение|протокол|регламент)\b"
    ),
    flags=re.IGNORECASE,
)


def extract_docx(
    body: bytes,
    *,
    content_type: str,
    final_url: str,
) -> ExtractedContent:
    try:
        from docx import Document

        document = Document(BytesIO(body))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            value = normalize_text(paragraph.text)

            if value:
                parts.append(value)

        for table in document.tables:
            for row in table.rows:
                values = [
                    compact_text(cell.text)
                    for cell in row.cells
                    if compact_text(cell.text)
                ]

                if values:
                    parts.append(" | ".join(values))

        text = normalize_text("\n".join(parts))

    except Exception as exc:
        raise CorruptDocumentError(
            f"Ошибка обработки DOCX: {exc}"
        ) from exc

    if len(text) < settings.min_text_length:
        raise EmptyContentError(
            "В DOCX не удалось извлечь достаточно текста"
        )

    # Реквизиты берутся из текста самого документа, как у PDF.
    # Раньше офисные файлы получали их только по наследству от
    # родительской HTML-карточки, а найденные прямо в /upload/
    # оставались вовсе без даты, номера и вида акта.
    title, legal_metadata = describe_document(text, final_url)

    return ExtractedContent(
        title=title,
        section="Документ DOCX",
        doc_type="file",
        content_type=(
            content_type
            or "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        pages=((1, text),),
        document_number=legal_metadata.document_number,
        document_date=legal_metadata.document_date,
        document_kind=legal_metadata.document_kind,
    )


def extract_xlsx(
    body: bytes,
    *,
    content_type: str,
    final_url: str,
) -> ExtractedContent:
    workbook: Any = None

    try:
        from openpyxl import load_workbook

        workbook = load_workbook(
            BytesIO(body),
            read_only=True,
            data_only=True,
        )

        pages: list[tuple[int, str]] = []

        for page_number, sheet in enumerate(
            workbook.worksheets,
            start=1,
        ):
            parts = [f"Лист: {sheet.title}"]

            for row in sheet.iter_rows(values_only=True):
                values = [
                    compact_text(str(cell))
                    for cell in row
                    if cell is not None
                    and compact_text(str(cell))
                ]

                if values:
                    parts.append(" | ".join(values))

            text = normalize_text("\n".join(parts))

            if text:
                pages.append((page_number, text))

    except Exception as exc:
        raise CorruptDocumentError(
            f"Ошибка обработки XLSX: {exc}"
        ) from exc

    finally:
        if workbook is not None:
            workbook.close()

    total_length = sum(len(text) for _, text in pages)

    if total_length < settings.min_text_length:
        raise EmptyContentError(
            "В XLSX не удалось извлечь достаточно текста"
        )

    # Реквизиты берутся из текста самого документа, как у PDF.
    # Раньше офисные файлы получали их только по наследству от
    # родительской HTML-карточки, а найденные прямо в /upload/
    # оставались вовсе без даты, номера и вида акта.
    title, legal_metadata = describe_document(pages[0][1] if pages else "", final_url)

    return ExtractedContent(
        title=title,
        section="Таблица XLSX",
        doc_type="file",
        content_type=(
            content_type
            or "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        pages=tuple(pages),
        document_number=legal_metadata.document_number,
        document_date=legal_metadata.document_date,
        document_kind=legal_metadata.document_kind,
    )


def decode_text_document(body: bytes) -> str:
    for encoding in (
        "utf-8-sig",
        "cp1251",
    ):
        try:
            return body.decode(encoding)
        except UnicodeDecodeError:
            continue

    return body.decode("utf-8", errors="replace")


def extract_text_document(
    body: bytes,
    *,
    content_type: str,
    final_url: str,
) -> ExtractedContent:
    text = normalize_text(
        decode_text_document(body)
    )

    if len(text) < settings.min_text_length:
        raise EmptyContentError(
            "В текстовом файле недостаточно текста"
        )

    # Реквизиты берутся из текста самого документа, как у PDF.
    # Раньше офисные файлы получали их только по наследству от
    # родительской HTML-карточки, а найденные прямо в /upload/
    # оставались вовсе без даты, номера и вида акта.
    title, legal_metadata = describe_document(text, final_url)

    return ExtractedContent(
        title=title,
        section="Текстовый файл",
        doc_type="file",
        content_type=content_type or "text/plain",
        pages=((1, text),),
        document_number=legal_metadata.document_number,
        document_date=legal_metadata.document_date,
        document_kind=legal_metadata.document_kind,
    )



# Управляющее слово RTF: обратный слэш, до 32 латинских букв,
# необязательный числовой параметр и один необязательный пробел-разделитель.
# `(?![a-zA-Z])` обязателен: без него `\par` съедал бы начало `\pard`
# и оставлял в тексте букву «d».
RTF_CONTROL_RE = re.compile(
    r"\\[a-zA-Z]{1,32}(?:-?\d{1,10})?[ ]?"
)
RTF_BREAK_RE = re.compile(
    r"\\(?:par|pard|line|sect|page|row|cell)"
    r"(?![a-zA-Z])[ ]?"
)
RTF_TAB_RE = re.compile(r"\\tab(?![a-zA-Z])[ ]?")
RTF_ESCAPED_SYMBOL_RE = re.compile(r"\\([{}\\])")
RTF_HEX_RE = re.compile(r"\\'([0-9a-fA-F]{2})")
RTF_UNICODE_RE = re.compile(r"\\u(-?\d+)\s*\??")

RTF_SKIP_GROUPS = (
    "fonttbl", "colortbl", "stylesheet", "info", "pict",
    "object", "themedata", "colorschememapping", "latentstyles",
    "datastore", "listtable", "listoverridetable", "rsidtbl",
    "generator", "filetbl", "xmlnstbl",
)


def extract_rtf(
    body: bytes,
    *,
    content_type: str,
    final_url: str,
) -> ExtractedContent:
    """
    Извлекает текст из RTF без внешних зависимостей.

    Разбор намеренно грубый: служебные группы отбрасываются целиком,
    управляющие слова удаляются, escape-последовательности
    раскрываются. Для поиска этого достаточно, а тянуть на общий сервер
    конвертер ради полутора тысяч файлов не нужно.
    """

    raw = body.decode("cp1251", errors="replace")

    # Отбрасываем служебные группы вместе с содержимым.
    for name in RTF_SKIP_GROUPS:
        marker = "{\\" + name
        while True:
            start = raw.find(marker)

            if start < 0:
                break

            depth = 0
            index = start

            while index < len(raw):
                if raw[index] == "{" and raw[index - 1: index] != "\\":
                    depth += 1
                elif raw[index] == "}" and raw[index - 1: index] != "\\":
                    depth -= 1

                    if depth == 0:
                        break

                index += 1

            raw = raw[:start] + raw[index + 1:]

    # Абзацы и переводы строк — до удаления остальных управляющих слов,
    # иначе границы предложений склеятся.
    raw = RTF_BREAK_RE.sub("\n", raw)
    raw = RTF_TAB_RE.sub(" ", raw)

    def replace_unicode(match: "re.Match[str]") -> str:
        code = int(match.group(1))

        if code < 0:
            code += 65536

        try:
            return chr(code)
        except ValueError:
            return ""

    raw = RTF_UNICODE_RE.sub(replace_unicode, raw)

    def replace_hex(match: "re.Match[str]") -> str:
        try:
            return bytes(
                [int(match.group(1), 16)]
            ).decode("cp1251", errors="replace")
        except ValueError:
            return ""

    raw = RTF_HEX_RE.sub(replace_hex, raw)

    # Экранированные фигурные скобки и слэш возвращаем в текст до того,
    # как удалим управляющие слова: иначе `\{` будет принят за одно из них.
    raw = RTF_ESCAPED_SYMBOL_RE.sub(
        lambda match: "\x00" + match.group(1),
        raw,
    )

    raw = RTF_CONTROL_RE.sub("", raw)
    raw = raw.replace("{", "").replace("}", "")
    raw = raw.replace("\x00", "")

    text = normalize_text(raw)

    if len(text) < settings.min_text_length:
        raise EmptyContentError(
            "RTF не содержит полезного текста"
        )

    # Реквизиты берутся из текста самого документа, как у PDF.
    # Раньше офисные файлы получали их только по наследству от
    # родительской HTML-карточки, а найденные прямо в /upload/
    # оставались вовсе без даты, номера и вида акта.
    title, legal_metadata = describe_document(text, final_url)

    return ExtractedContent(
        title=title,
        section="Документ RTF",
        doc_type="file",
        content_type=content_type or "application/rtf",
        pages=((1, text),),
        extraction_version="rtf-v1",
        document_number=legal_metadata.document_number,
        document_date=legal_metadata.document_date,
        document_kind=legal_metadata.document_kind,
    )


def extract_opendocument(
    body: bytes,
    *,
    content_type: str,
    final_url: str,
) -> ExtractedContent:
    """
    Извлекает текст из ODT и ODS.

    OpenDocument — это ZIP с XML внутри, поэтому достаточно стандартной
    библиотеки.
    """

    import xml.etree.ElementTree as ElementTree
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(body)) as archive:
            names = set(archive.namelist())

            if "content.xml" not in names:
                raise CorruptDocumentError(
                    "В OpenDocument нет content.xml"
                )

            payload = archive.read("content.xml")

    except zipfile.BadZipFile as exc:
        raise CorruptDocumentError(
            f"Повреждён контейнер OpenDocument: {exc}"
        ) from exc

    try:
        root = ElementTree.fromstring(payload)
    except ElementTree.ParseError as exc:
        raise CorruptDocumentError(
            f"Повреждён content.xml: {exc}"
        ) from exc

    parts: list[str] = []

    for element in root.iter():
        tag = element.tag.rsplit("}", 1)[-1]

        if element.text and tag not in {"annotation", "note"}:
            parts.append(element.text)

        # Разделители строк и ячеек, чтобы текст не слипался.
        if tag in {"p", "h", "table-row", "list-item"}:
            parts.append("\n")
        elif tag in {"table-cell", "tab"}:
            parts.append("\t")

        if element.tail:
            parts.append(element.tail)

    text = normalize_text("".join(parts))

    if len(text) < settings.min_text_length:
        raise EmptyContentError(
            "OpenDocument не содержит полезного текста"
        )

    suffix = PurePosixPath(
        urlparse(final_url).path.lower()
    ).suffix

    # Реквизиты берутся из текста самого документа, как у PDF.
    # Раньше офисные файлы получали их только по наследству от
    # родительской HTML-карточки, а найденные прямо в /upload/
    # оставались вовсе без даты, номера и вида акта.
    title, legal_metadata = describe_document(text, final_url)

    return ExtractedContent(
        title=title,
        section=(
            "Таблица OpenDocument"
            if suffix == ".ods"
            else "Документ OpenDocument"
        ),
        doc_type="file",
        content_type=(
            content_type
            or "application/vnd.oasis.opendocument.text"
        ),
        pages=((1, text),),
        extraction_version="odf-v1",
        document_number=legal_metadata.document_number,
        document_date=legal_metadata.document_date,
        document_kind=legal_metadata.document_kind,
    )


def extract_content(
    body: bytes | Path,
    *,
    content_type: str,
    final_url: str,
    pdf_ocr_mode: str = PDF_OCR_MODE_BOUNDED,
) -> ExtractedContent:
    lowered_type = (
        content_type
        or ""
    ).lower()

    path = urlparse(
        final_url
    ).path.lower()

    suffix = PurePosixPath(path).suffix

    is_pdf = (
        isinstance(body, Path)
        or (
            isinstance(body, bytes)
            and body.startswith(b"%PDF-")
        )
        or lowered_type == "application/pdf"
        or path.endswith(".pdf")
    )

    if is_pdf:
        return extract_pdf(
            body,
            content_type=content_type,
            final_url=final_url,
            ocr_mode=pdf_ocr_mode,
        )

    if isinstance(body, Path):
        raise UnsupportedContentTypeError(
            "Временный файл создан не для PDF"
        )

    is_docx = (
        suffix == ".docx"
        or "wordprocessingml.document" in lowered_type
    )

    if is_docx:
        return extract_docx(
            body,
            content_type=content_type,
            final_url=final_url,
        )

    is_xlsx = (
        suffix == ".xlsx"
        or "spreadsheetml.sheet" in lowered_type
    )

    if is_xlsx:
        return extract_xlsx(
            body,
            content_type=content_type,
            final_url=final_url,
        )

    if suffix == ".rtf" or "rtf" in lowered_type:
        return extract_rtf(
            body,
            content_type=content_type,
            final_url=final_url,
        )

    if (
        suffix in {".odt", ".ods"}
        or "opendocument" in lowered_type
    ):
        return extract_opendocument(
            body,
            content_type=content_type,
            final_url=final_url,
        )

    is_html = (
        "html" in lowered_type
        or body.lstrip().startswith(
            (
                b"<!DOCTYPE html",
                b"<!doctype html",
                b"<html",
                b"<HTML",
            )
        )
    )

    if is_html:
        return extract_html(
            body,
            content_type=content_type,
            final_url=final_url,
        )

    if (
        suffix in {".csv", ".txt"}
        or lowered_type in {"text/csv", "text/plain"}
    ):
        return extract_text_document(
            body,
            content_type=content_type,
            final_url=final_url,
        )

    raise UnsupportedContentTypeError(
        "Неподдерживаемый тип содержимого: "
        f"{content_type or 'unknown'}"
    )


def calculate_content_hash(
    content: ExtractedContent,
) -> str:
    metadata_lines = [
        f"TITLE:{content.title}",
        f"SECTION:{content.section}",
        (
            "CONTENT_CATEGORY:"
            f"{content.content_category or ''}"
        ),
        (
            "SITE_SECTION:"
            f"{content.site_section or ''}"
        ),
        (
            "PUBLISHED_AT:"
            f"{content.published_at or ''}"
        ),
        (
            "DOCUMENT_NUMBER:"
            f"{content.document_number or ''}"
        ),
        (
            "DOCUMENT_DATE:"
            f"{content.document_date or ''}"
        ),
        (
            "DOCUMENT_KIND:"
            f"{content.document_kind or ''}"
        ),
    ]

    if content.extraction_version:
        metadata_lines.extend(
            [
                (
                    "EXTRACTION_VERSION:"
                    f"{content.extraction_version}"
                ),
                (
                    "OCR_PAGES:"
                    + ",".join(
                        str(value)
                        for value in content.ocr_pages
                    )
                ),
                (
                    "OCR_ATTEMPTED_PAGES:"
                    + ",".join(
                        str(value)
                        for value in content.ocr_attempted_pages
                    )
                ),
                (
                    "OCR_REJECTED_PAGES:"
                    + ",".join(
                        str(value)
                        for value in content.ocr_rejected_pages
                    )
                ),
                (
                    "OCR_TIMED_OUT_PAGES:"
                    + ",".join(
                        str(value)
                        for value in content.ocr_timed_out_pages
                    )
                ),
                (
                    "OCR_BUDGET_SKIPPED_PAGES:"
                    + ",".join(
                        str(value)
                        for value in content.ocr_budget_skipped_pages
                    )
                ),
                (
                    "OCR_DEFERRED_PAGES:"
                    + ",".join(
                        str(value)
                        for value in content.ocr_deferred_pages
                    )
                ),
                (
                    "UNREADABLE_PAGES:"
                    + ",".join(
                        str(value)
                        for value in content.unreadable_pages
                    )
                ),
            ]
        )

    for attachment in sorted(
        content.attachments,
        key=lambda item: item.url,
    ):
        metadata_lines.append(
            "ATTACHMENT:"
            f"{attachment.url}|"
            f"{attachment.title}"
        )

    body_text = "\n\f\n".join(
        f"PAGE:{page_number}\n{text}"
        for page_number, text in content.pages
    )

    hash_input = (
        "\n".join(metadata_lines)
        + "\n\f\n"
        + body_text
    )

    return hashlib.sha256(
        hash_input.encode(
            "utf-8"
        )
    ).hexdigest()


def find_chunk_end(
    text: str,
    *,
    start: int,
    proposed_end: int,
) -> int:
    minimum_end = min(
        start + 500,
        len(text),
    )

    if proposed_end >= len(text):
        return len(text)

    search_start = max(
        minimum_end,
        proposed_end - 700,
    )

    candidates: list[int] = []

    for separator in (
        "\n\n",
        ". ",
        "! ",
        "? ",
        "; ",
        "\n",
    ):
        position = text.rfind(
            separator,
            search_start,
            proposed_end,
        )

        if position >= minimum_end:
            candidates.append(
                position
                + len(separator)
            )

    if candidates:
        return max(candidates)

    return proposed_end


def split_into_chunks(
    content: ExtractedContent,
) -> list[TextChunk]:
    result: list[TextChunk] = []
    global_chunk_number = 0

    for page_number, page_text in (
        content.pages
    ):
        normalized = normalize_text(
            page_text
        )

        if not normalized:
            continue

        start = 0

        while start < len(normalized):
            proposed_end = min(
                start + settings.chunk_size,
                len(normalized),
            )

            end = find_chunk_end(
                normalized,
                start=start,
                proposed_end=proposed_end,
            )

            if end <= start:
                end = proposed_end

            chunk_text = normalized[
                start:end
            ].strip()

            if chunk_text:
                global_chunk_number += 1

                result.append(
                    TextChunk(
                        page_number=page_number,
                        chunk_number=(
                            global_chunk_number
                        ),
                        text=chunk_text,
                    )
                )

            if end >= len(normalized):
                break

            next_start = max(
                end - settings.chunk_overlap,
                start + 1,
            )

            start = next_start

    if not result:
        raise EmptyContentError(
            "После разбиения не осталось "
            "текстовых фрагментов"
        )

    return result
