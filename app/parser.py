from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin

import fitz
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from openpyxl import load_workbook

from .settings import CHUNK_OVERLAP, CHUNK_SIZE


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = " ".join(text.split())
    return text.strip()


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    step = max(1, size - overlap)

    while start < len(text):
        chunk = text[start:start + size].strip()
        if chunk:
            chunks.append(chunk)
        start += step

    return chunks


def allowed_pdf_link(href: str) -> bool:
    href = href.lower()
    return ".pdf" in href or href.endswith(".pdf")


def extract_html_page(html: str, base_url: str) -> dict:
    soup = BeautifulSoup(html, "lxml")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript", "form"]):
        tag.decompose()

    title = ""
    if soup.title and soup.title.string:
        title = normalize_text(soup.title.string)

    if not title:
        h1 = soup.find("h1")
        if h1:
            title = normalize_text(h1.get_text(" ", strip=True))

    if not title:
        meta = soup.find("meta", attrs={"property": "og:title"})
        if meta and meta.get("content"):
            title = normalize_text(meta["content"])

    if not title:
        title = base_url

    pdf_links: list[str] = []
    for a in soup.find_all("a", href=True):
        href = urljoin(base_url, a["href"])
        if allowed_pdf_link(href):
            pdf_links.append(href.split("#")[0])

    text = soup.get_text(" ", strip=True)
    text = normalize_text(text)

    return {
        "title": title,
        "text": text,
        "pdf_links": sorted(set(pdf_links)),
    }


def extract_pdf_pages(pdf_bytes: bytes) -> list[dict]:
    pages: list[dict] = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page_no in range(len(doc)):
            page = doc[page_no]
            text = normalize_text(page.get_text("text"))
            if text:
                pages.append({
                    "page_number": page_no + 1,
                    "text": text,
                })
    return pages


def extract_docx_text(data: bytes) -> list[dict]:
    doc = DocxDocument(BytesIO(data))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_values: list[str] = []
            for cell in row.cells:
                value = normalize_text(cell.text)
                if value:
                    row_values.append(value)
            if row_values:
                parts.append(" | ".join(row_values))

    text = normalize_text("\n".join(parts))
    return [{"page_number": 1, "text": text}] if text else []


def extract_xlsx_text(data: bytes) -> list[dict]:
    wb = load_workbook(BytesIO(data), data_only=True)
    pages: list[dict] = []

    for sheet in wb.worksheets:
        parts: list[str] = [f"Лист: {sheet.title}"]
        for row in sheet.iter_rows(values_only=True):
            values = []
            for cell in row:
                if cell is None:
                    continue
                value = normalize_text(str(cell))
                if value:
                    values.append(value)
            if values:
                parts.append(" | ".join(values))

        text = normalize_text("\n".join(parts))
        if text:
            pages.append({
                "page_number": 1,
                "text": text,
                "section": sheet.title,
            })

    return pages


def extract_plain_text(data: bytes) -> list[dict]:
    text = normalize_text(data.decode("utf-8", errors="ignore"))
    return [{"page_number": 1, "text": text}] if text else []


def extract_document_bytes(filename: str, content_type: str, data: bytes) -> list[dict]:
    filename = filename.lower()
    content_type = (content_type or "").lower()

    if filename.endswith(".pdf") or "pdf" in content_type:
        return extract_pdf_pages(data)

    if filename.endswith(".docx") or "officedocument.wordprocessingml.document" in content_type:
        return extract_docx_text(data)

    if filename.endswith(".xlsx") or "spreadsheetml.sheet" in content_type:
        return extract_xlsx_text(data)

    if filename.endswith(".txt") or filename.endswith(".md") or "text/" in content_type:
        return extract_plain_text(data)

    if filename.endswith(".html") or filename.endswith(".htm") or "html" in content_type:
        html = data.decode("utf-8", errors="ignore")
        extracted = extract_html_page(html, filename)
        return [{"page_number": 1, "text": extracted["text"]}] if extracted["text"] else []

    return []
